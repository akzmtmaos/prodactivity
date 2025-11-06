# notes/views.py
from rest_framework import generics, status
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.db.models import Q
from .models import Notebook, Note
from .serializers import NotebookSerializer, NoteSerializer
import os
from django.conf import settings
from rest_framework.decorators import api_view, permission_classes
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import docx
import tempfile
from django.utils import timezone
import io
import re
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import base64
from PIL import Image as PILImage
from io import BytesIO

class NotebookListCreateView(generics.ListCreateAPIView):
    serializer_class = NotebookSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        # Filter by archive status and delete status
        is_archived = self.request.query_params.get('archived', 'false').lower() == 'true'
        is_deleted = self.request.query_params.get('is_deleted', 'false').lower() == 'true'
        return Notebook.objects.filter(user=self.request.user, is_archived=is_archived, is_deleted=is_deleted)

class NotebookRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    serializer_class = NotebookSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Notebook.objects.filter(user=self.request.user)
    
    def destroy(self, request, *args, **kwargs):
        notebook = self.get_object()
        if notebook.is_deleted:
            # Hard delete if already soft deleted
            notebook.notes.update(is_deleted=True, deleted_at=timezone.now())
            notebook.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        # Soft delete the notebook
        notebook.is_deleted = True
        notebook.deleted_at = timezone.now()
        # Soft delete all notes in this notebook
        notebook.notes.update(is_deleted=True, deleted_at=timezone.now())
        notebook.save()
        return Response(status=status.HTTP_204_NO_CONTENT)

    def partial_update(self, request, *args, **kwargs):
        notebook = self.get_object()
        is_archived = request.data.get('is_archived', None)
        is_deleted = request.data.get('is_deleted', None)
        
        if is_archived is not None:
            notebook.is_archived = is_archived
            notebook.archived_at = timezone.now() if is_archived else None
            notebook.save()
        
        if is_deleted is not None:
            notebook.is_deleted = is_deleted
            notebook.deleted_at = None if not is_deleted else timezone.now()
            # Also restore/delete notes in the notebook
            notebook.notes.update(is_deleted=is_deleted, deleted_at=timezone.now() if is_deleted else None)
            notebook.save()
        
        serializer = self.get_serializer(notebook)
        return Response(serializer.data)

class NoteListCreateView(generics.ListCreateAPIView):
    serializer_class = NoteSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        # Filter by archive status
        is_archived = self.request.query_params.get('archived', 'false').lower() == 'true'
        queryset = Note.objects.filter(user=self.request.user, is_deleted=False, is_archived=is_archived)
        
        # Filter by notebook if provided (unless global search is requested)
        notebook_id = self.request.query_params.get('notebook', None)
        global_search = self.request.query_params.get('global_search', 'false').lower() == 'true'
        
        if notebook_id is not None and not global_search:
            queryset = queryset.filter(notebook_id=notebook_id)
        
        # Search functionality
        search = self.request.query_params.get('search', None)
        if search is not None:
            queryset = queryset.filter(
                Q(title__icontains=search) | Q(content__icontains=search)
            )
        
        return queryset

class NoteRetrieveUpdateDestroyView(generics.RetrieveUpdateDestroyAPIView):
    serializer_class = NoteSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Note.objects.filter(user=self.request.user)
    
    def destroy(self, request, *args, **kwargs):
        note = self.get_object()
        if note.is_deleted:
            note.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        note.is_deleted = True
        note.deleted_at = timezone.now()
        note.save()
        return Response(status=status.HTTP_204_NO_CONTENT)

    def update(self, request, *args, **kwargs):
        note = self.get_object()
        print(f"[DEBUG] PUT /api/notes/{note.id}/ - data: {request.data}")
        
        # Merge request data with existing note data to ensure all fields are present
        # This allows partial updates while maintaining full object validation
        update_data = {
            'title': request.data.get('title', note.title),
            'content': request.data.get('content', note.content),
            'notebook': request.data.get('notebook', note.notebook.id),
            'note_type': request.data.get('note_type', note.note_type),
            'priority': request.data.get('priority', note.priority),
            'is_urgent': request.data.get('is_urgent', note.is_urgent),
            'tags': request.data.get('tags', note.tags or ''),
            'is_deleted': request.data.get('is_deleted', note.is_deleted),
            'is_archived': request.data.get('is_archived', note.is_archived),
        }
        
        # Update the note with the merged data
        serializer = self.get_serializer(note, data=update_data, partial=False)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        else:
            print(f"[DEBUG] Validation errors: {serializer.errors}")
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def partial_update(self, request, *args, **kwargs):
        note = self.get_object()
        print(f"[DEBUG] PATCH /api/notes/{note.id}/ - data: {request.data}")
        
        # Use serializer with partial=True to handle all fields properly
        serializer = self.get_serializer(note, data=request.data, partial=True)
        if serializer.is_valid():
            # Save the note with all updates from serializer
            serializer.save()
            
            # Handle special timestamp fields that need custom logic
            needs_save = False
            if 'is_archived' in request.data:
                is_archived = request.data.get('is_archived')
                note.archived_at = timezone.now() if is_archived else None
                needs_save = True
                print(f"[DEBUG] Note {note.id} updated: is_archived={note.is_archived}, archived_at={note.archived_at}")
            
            if 'is_deleted' in request.data:
                is_deleted = request.data.get('is_deleted')
                note.deleted_at = None if not is_deleted else timezone.now()
                needs_save = True
                print(f"[DEBUG] Note {note.id} updated: is_deleted={note.is_deleted}, deleted_at={note.deleted_at}")
            
            # Save timestamp fields if they were updated
            if needs_save:
                note.save(update_fields=['archived_at', 'deleted_at'])
            
            # Return updated note data
            updated_serializer = self.get_serializer(note)
            return Response(updated_serializer.data)
        else:
            print(f"[DEBUG] Validation errors: {serializer.errors}")
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

def convert_docx_to_html(doc):
    """Convert DOCX document to HTML while preserving formatting like bullet lists and tables"""
    html_parts = []
    
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            paragraph = None
            for p in doc.paragraphs:
                if p._element == element:
                    paragraph = p
                    break
            
            if paragraph:
                # Check if it's a list item
                if paragraph.style.name.startswith('List'):
                    # Extract list level and create appropriate HTML
                    list_level = 0
                    if hasattr(paragraph.style, 'name') and 'List' in paragraph.style.name:
                        try:
                            list_level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 0
                        except:
                            list_level = 0
                    
                    # Create list item
                    indent = '  ' * list_level
                    html_parts.append(f'{indent}<li>{paragraph.text}</li>')
                else:
                    # Regular paragraph
                    if paragraph.text.strip():
                        html_parts.append(f'<p>{paragraph.text}</p>')
        
        elif element.tag.endswith('tbl'):  # Table
            table_html = convert_table_to_html(element, doc)
            if table_html:
                html_parts.append(table_html)
    
    # Post-process to create proper list structure
    final_html = []
    in_list = False
    list_level = 0
    
    for line in html_parts:
        if line.strip().startswith('<li>'):
            if not in_list:
                final_html.append('<ul>')
                in_list = True
            final_html.append(line)
        else:
            if in_list:
                final_html.append('</ul>')
                in_list = False
            final_html.append(line)
    
    if in_list:
        final_html.append('</ul>')
    
    return '\n'.join(final_html)

def convert_table_to_html(table_element, doc):
    """Convert DOCX table to HTML"""
    try:
        html_parts = ['<table class="notion-table">']
        
        # Find the table in the document
        table = None
        for t in doc.tables:
            if t._element == table_element:
                table = t
                break
        
        if not table:
            return ""
        
        # Add table rows
        for i, row in enumerate(table.rows):
            if i == 0:  # Header row
                html_parts.append('<thead><tr>')
                for cell in row.cells:
                    html_parts.append(f'<th>{cell.text}</th>')
                html_parts.append('</tr></thead><tbody>')
            else:  # Data rows
                html_parts.append('<tr>')
                for cell in row.cells:
                    html_parts.append(f'<td>{cell.text}</td>')
                html_parts.append('</tr>')
        
        html_parts.append('</tbody></table>')
        return '\n'.join(html_parts)
    except Exception as e:
        return f"<p>Table conversion error: {str(e)}</p>"

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def convert_doc(request):
    if 'file' not in request.FILES:
        return Response({'error': 'No file provided'}, status=status.HTTP_400_BAD_REQUEST)
    
    file = request.FILES['file']
    if not file.name.endswith(('.doc', '.docx')):
        return Response({'error': 'Invalid file type. Please upload a DOC or DOCX file.'}, 
                       status=status.HTTP_400_BAD_REQUEST)
    
    try:
        # Save the uploaded file temporarily
        temp_path = default_storage.save(f'temp/{file.name}', ContentFile(file.read()))
        temp_file_path = os.path.join(settings.MEDIA_ROOT, temp_path)
        
        # Open and read the document
        doc = docx.Document(temp_file_path)
        
        # Extract content with formatting preserved
        html_content = convert_docx_to_html(doc)
        
        # Clean up the temporary file
        default_storage.delete(temp_path)
        
        return Response({'text': html_content})
    except Exception as e:
        # Clean up in case of error
        if 'temp_path' in locals():
            default_storage.delete(temp_path)
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def deleted_notes(request):
    notes = Note.objects.filter(user=request.user, is_deleted=True)
    print(f"[DEBUG] Trash API - User: {request.user}, Deleted Notes: {list(notes.values('id', 'title', 'is_deleted', 'deleted_at'))}")
    serializer = NoteSerializer(notes, many=True)
    return Response(serializer.data)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def archived_notes(request):
    notes = Note.objects.filter(user=request.user, is_archived=True, is_deleted=False)
    serializer = NoteSerializer(notes, many=True)
    return Response(serializer.data)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def archived_notebooks(request):
    notebooks = Notebook.objects.filter(user=request.user, is_archived=True)
    serializer = NotebookSerializer(notebooks, many=True)
    return Response(serializer.data) 

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def global_search_notes(request):
    """
    Global search endpoint to search across all notes and notebooks
    """
    search_query = request.query_params.get('q', '')
    if not search_query.strip():
        return Response({'results': []})
    
    results = []
    
    # Search across all notes (not deleted, not archived)
    notes = Note.objects.filter(
        user=request.user,
        is_deleted=False,
        is_archived=False
    ).filter(
        Q(title__icontains=search_query) | 
        Q(content__icontains=search_query)
    ).select_related('notebook').order_by('-updated_at')
    
    # Add notes to results
    for note in notes:
        results.append({
            'id': note.id,
            'type': 'note',
            'title': note.title,
            'content': note.content[:200] + '...' if len(note.content) > 200 else note.content,
            'notebook_id': note.notebook.id,
            'notebook_name': note.notebook.name,
            'created_at': note.created_at,
            'updated_at': note.updated_at,
            'url': f'/notes/{note.id}'
        })
    
    # Search across all notebooks (not archived)
    notebooks = Notebook.objects.filter(
        user=request.user,
        is_archived=False
    ).filter(
        Q(name__icontains=search_query)
    ).order_by('-updated_at')
    
    # Add notebooks to results
    for notebook in notebooks:
        results.append({
            'id': notebook.id,
            'type': 'notebook',
            'title': notebook.name,
            'content': f'Notebook with {notebook.notes_count} notes',
            'notebook_id': notebook.id,
            'notebook_name': notebook.name,
            'created_at': notebook.created_at,
            'updated_at': notebook.updated_at,
            'url': f'/notes?notebook={notebook.id}'
        })
    
    # Sort all results by updated_at (most recent first)
    results.sort(key=lambda x: x['updated_at'], reverse=True)
    
    return Response({'results': results})

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def urgent_notes_and_notebooks(request):
    """Get all urgent notes and notebooks that need immediate attention"""
    try:
        # Get urgent notes
        urgent_notes = Note.objects.filter(
            user=request.user,
            is_deleted=False,
            is_archived=False
        ).filter(
            Q(is_urgent=True) | Q(priority='urgent') | Q(priority='high')
        ).select_related('notebook').order_by('-updated_at')
        
        # Get urgent notebooks
        # Get urgent notes only (notebook urgency_level field removed)
        urgent_notebooks = Notebook.objects.none()  # No longer filtering by urgency_level
        
        # Get notes from urgent notebooks (notebook urgency_level field removed)
        notes_from_urgent_notebooks = Note.objects.none()  # No longer filtering by urgency_level
        
        # Combine and deduplicate
        all_urgent_notes = list(urgent_notes) + list(notes_from_urgent_notebooks)
        unique_urgent_notes = list({note.id: note for note in all_urgent_notes}.values())
        
        # Serialize the data
        note_serializer = NoteSerializer(unique_urgent_notes, many=True)
        notebook_serializer = NotebookSerializer(urgent_notebooks, many=True)
        
        return Response({
            'urgent_notes': note_serializer.data,
            'urgent_notebooks': notebook_serializer.data,
            'total_urgent_items': len(unique_urgent_notes) + len(urgent_notebooks)
        })
        
    except Exception as e:
        return Response(
            {'error': f'Failed to fetch urgent items: {str(e)}'}, 
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def notes_by_type(request):
    """Get notes filtered by type (study, meeting, etc.)"""
    note_type = request.query_params.get('type', None)
    if not note_type:
        return Response({'error': 'Note type parameter is required'}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        notes = Note.objects.filter(
            user=request.user,
            is_deleted=False,
            is_archived=False,
            note_type=note_type
        ).select_related('notebook').order_by('-updated_at')
        
        serializer = NoteSerializer(notes, many=True)
        return Response({
            'notes': serializer.data,
            'type': note_type,
            'count': len(notes)
        })
        
    except Exception as e:
        return Response(
            {'error': f'Failed to fetch notes by type: {str(e)}'}, 
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def notebooks_by_type(request):
    """Get notebooks filtered by type (DEPRECATED: notebook_type field removed)"""
    # Notebook type field has been removed, return all notebooks
    try:
        notebooks = Notebook.objects.filter(
            user=request.user,
            is_archived=False
        ).order_by('-updated_at')
        
        serializer = NotebookSerializer(notebooks, many=True)
        return Response({
            'notebooks': serializer.data,
            'count': len(notebooks),
            'message': 'Notebook type filtering has been removed. Returning all notebooks.'
        })
        
    except Exception as e:
        return Response(
            {'error': f'Failed to fetch notebooks: {str(e)}'}, 
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

def clean_html_content(html_content):
    """Clean HTML content for export by removing tags and entities"""
    if not html_content:
        return ""
    
    # Remove HTML tags
    content = re.sub(r'<[^>]+>', '', html_content)
    
    # Replace HTML entities
    content = content.replace('&nbsp;', ' ')
    content = content.replace('&amp;', '&')
    content = content.replace('&lt;', '<')
    content = content.replace('&gt;', '>')
    content = content.replace('&quot;', '"')
    content = content.replace('&#39;', "'")
    
    # Clean up extra whitespace and line breaks
    content = re.sub(r'\n\s*\n', '\n\n', content)  # Remove multiple empty lines
    content = re.sub(r' +', ' ', content)  # Remove multiple spaces
    content = content.strip()  # Remove leading/trailing whitespace
    
    return content

def extract_images_from_html(html_content):
    """Extract images from HTML content and return list of image data"""
    if not html_content:
        return []
    
    images = []
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for img_tag in soup.find_all('img'):
        src = img_tag.get('src', '')
        if not src:
            continue
        
        try:
            # Handle base64 data URLs
            if src.startswith('data:image'):
                # Extract base64 data
                header, data = src.split(',', 1)
                image_format = header.split('/')[1].split(';')[0]  # e.g., 'png', 'jpeg'
                
                # Decode base64
                image_data = base64.b64decode(data)
                
                # Open image with PIL to validate and get dimensions
                img = PILImage.open(BytesIO(image_data))
                
                images.append({
                    'data': image_data,
                    'format': image_format,
                    'width': img.width,
                    'height': img.height
                })
            # Handle regular URLs (could be extended to download images)
            elif src.startswith('http://') or src.startswith('https://'):
                # For now, skip external URLs (could be extended to download)
                print(f"Skipping external image URL: {src}")
                continue
        except Exception as e:
            print(f"Error processing image: {str(e)}")
            continue
    
    return images

def process_html_content_with_images(html_content):
    """Process HTML content and extract text and images separately"""
    if not html_content:
        return {'text': '', 'images': []}
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Extract images
    images = extract_images_from_html(html_content)
    
    # Replace img tags with placeholders and extract text
    text_parts = []
    img_index = 0
    
    def process_element(element):
        nonlocal img_index
        if element.name == 'img':
            # Replace image with placeholder
            text_parts.append(f'\n[IMAGE_{img_index}]\n')
            img_index += 1
        elif hasattr(element, 'children'):
            # Process children recursively
            for child in element.children:
                if isinstance(child, str):
                    if child.strip():
                        text_parts.append(child.strip())
                else:
                    process_element(child)
        elif hasattr(element, 'string') and element.string:
            text = element.string.strip()
            if text:
                text_parts.append(text)
    
    # Process all elements
    for element in soup.children:
        if isinstance(element, str):
            if element.strip():
                text_parts.append(element.strip())
        else:
            process_element(element)
    
    # Get clean text
    clean_text = '\n'.join(text_parts)
    # Clean up HTML entities
    clean_text = clean_text.replace('&nbsp;', ' ')
    clean_text = clean_text.replace('&amp;', '&')
    clean_text = clean_text.replace('&lt;', '<')
    clean_text = clean_text.replace('&gt;', '>')
    clean_text = clean_text.replace('&quot;', '"')
    clean_text = clean_text.replace('&#39;', "'")
    # Normalize whitespace but preserve line breaks
    clean_text = re.sub(r'[ \t]+', ' ', clean_text)  # Multiple spaces to single space
    clean_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', clean_text)  # Multiple newlines to double newline
    
    return {'text': clean_text, 'images': images}

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def export_notes(request):
    """Export notes to PDF or DOC format"""
    try:
        note_id = request.data.get('note_id')
        format_type = request.data.get('format', 'pdf')
        
        if not note_id:
            return Response({'error': 'Note ID is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get the note
        try:
            note = Note.objects.get(id=note_id, user=request.user, is_deleted=False)
        except Note.DoesNotExist:
            return Response({'error': 'Note not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Process HTML content to extract text and images
        processed = process_html_content_with_images(note.content)
        clean_content = processed['text']
        images = processed['images']
        
        if format_type == 'pdf':
            # Create PDF
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                alignment=TA_LEFT
            )
            
            # Add title
            story.append(Paragraph(note.title, title_style))
            story.append(Spacer(1, 20))
            
            # Process content and images
            content_lines = clean_content.split('\n')
            image_index = 0
            
            for line in content_lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                elif line.startswith('[IMAGE_') and line.endswith(']'):
                    # Insert image
                    if image_index < len(images):
                        img_data = images[image_index]
                        try:
                            # Create temporary file for image
                            img_buffer = BytesIO(img_data['data'])
                            img = PILImage.open(img_buffer)
                            
                            # Calculate size (max width 5 inches, maintain aspect ratio)
                            max_width = 5 * inch
                            # Convert pixels to points (assuming 72 DPI)
                            img_width_pt = img.width * (72 / 96)  # Convert from 96 DPI to 72 DPI
                            img_height_pt = img.height * (72 / 96)
                            
                            if img_width_pt > max_width:
                                aspect_ratio = img_height_pt / img_width_pt
                                width = max_width
                                height = width * aspect_ratio
                            else:
                                width = img_width_pt
                                height = img_height_pt
                            
                            # Save image to temporary buffer
                            temp_img_buffer = BytesIO()
                            img.save(temp_img_buffer, format=img_data['format'].upper())
                            temp_img_buffer.seek(0)
                            
                            # Add image to PDF
                            pdf_image = ReportLabImage(temp_img_buffer, width=width, height=height)
                            story.append(pdf_image)
                            story.append(Spacer(1, 12))
                            image_index += 1
                        except Exception as e:
                            print(f"Error adding image to PDF: {str(e)}")
                            # Skip image if there's an error
                            image_index += 1
                else:
                    story.append(Paragraph(line, content_style))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            # Create response
            from django.http import HttpResponse
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{note.title}_export.pdf"'
            return response
            
        elif format_type == 'doc':
            # Create DOCX
            doc = Document()
            
            # Add title
            doc.add_heading(note.title, 0)
            
            # Process content and images
            content_lines = clean_content.split('\n')
            image_index = 0
            
            for line in content_lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()  # Empty paragraph for spacing
                elif line.startswith('[IMAGE_') and line.endswith(']'):
                    # Insert image
                    if image_index < len(images):
                        img_data = images[image_index]
                        try:
                            # Create temporary file for image
                            img_buffer = BytesIO(img_data['data'])
                            
                            # Add image to document (max width 5 inches)
                            paragraph = doc.add_paragraph()
                            run = paragraph.add_run()
                            
                            # Calculate size (max width 5 inches, maintain aspect ratio)
                            max_width = Inches(5)
                            # Convert pixels to inches (assuming 96 DPI)
                            img_width_in = img_data['width'] / 96
                            img_height_in = img_data['height'] / 96
                            
                            if img_width_in > 5:
                                aspect_ratio = img_height_in / img_width_in
                                width = max_width
                            else:
                                width = Inches(img_width_in)
                            
                            # Add picture
                            run.add_picture(img_buffer, width=width)
                            image_index += 1
                        except Exception as e:
                            print(f"Error adding image to DOCX: {str(e)}")
                            # Skip image if there's an error
                            image_index += 1
                else:
                    doc.add_paragraph(line)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Create response
            from django.http import HttpResponse
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{note.title}_export.docx"'
            return response
            
        else:
            return Response({'error': 'Unsupported format. Use "pdf" or "doc"'}, status=status.HTTP_400_BAD_REQUEST)
            
    except Exception as e:
        print(f"Export error: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': f'Export failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR) 