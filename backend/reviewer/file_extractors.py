"""
File text extraction utilities for various document formats.
Supports: PDF, DOCX, TXT, and other plain text formats.
"""
import logging
from typing import Tuple
from docx import Document
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> Tuple[str, str]:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num}: {e}")
                continue
        
        if not text_parts:
            return "", "No text could be extracted from the PDF"
        
        extracted_text = "\n\n".join(text_parts)
        return extracted_text.strip(), ""
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return "", f"Failed to read PDF file: {str(e)}"


def extract_text_from_docx(file_path: str) -> Tuple[str, str]:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        if not text_parts:
            return "", "No text could be extracted from the document"
        
        extracted_text = "\n\n".join(text_parts)
        return extracted_text.strip(), ""
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return "", f"Failed to read DOCX file: {str(e)}"


def extract_text_from_txt(file_path: str) -> Tuple[str, str]:
    """
    Extract text from a plain text file.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    if text.strip():
                        return text.strip(), ""
            except UnicodeDecodeError:
                continue
        
        return "", "Could not decode text file with any supported encoding"
        
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        return "", f"Failed to read text file: {str(e)}"


def extract_text_from_file(file_path: str, file_extension: str) -> Tuple[str, str]:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file
        file_extension: File extension (e.g., '.pdf', '.docx', '.txt')
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    file_extension = file_extension.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif file_extension in ['.txt', '.md', '.markdown', '.rst']:
        return extract_text_from_txt(file_path)
    else:
        return "", f"Unsupported file format: {file_extension}. Supported formats: PDF, DOCX, TXT, MD"

